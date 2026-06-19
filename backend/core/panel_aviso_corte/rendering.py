"""Renderizado de :class:`Panel` a PDF mediante Jinja2 + WeasyPrint."""

from __future__ import annotations

import base64
import contextlib
import logging
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import TYPE_CHECKING, Any

from jinja2 import Environment, FileSystemLoader, select_autoescape

from .errors import RenderingError

if TYPE_CHECKING:
    from .models import ExportMode, Panel, PanelImageRef

logger = logging.getLogger(__name__)

DOC_FONT = "Aptos"
TEAL_RGB = (0x3B, 0xA9, 0xAF)
TABLE_WIDTH_CM = 18.45
COL_WIDTHS_CM = (4.24, 4.98, 3.47, 5.76)
ROW_HEIGHTS_CM = {
    "title": 1.05,
    "meta_a": 0.73,
    "meta_b": 0.77,
    "meta_c": 0.96,
    "section": 0.53,
    "image": 9.82,
    "caption": 1.40,
    "caption_last": 1.43,
}
CELL_MARGIN_TWIPS = 104
PHOTO_WIDTH_CM = 7.36
PHOTO_HEIGHT_CM = ROW_HEIGHTS_CM["image"]
LOGO_WIDTH_CM = 5.49


def _template_dir() -> Path:
    bundled = Path(__file__).resolve().parent.parent.parent / "templates"
    if bundled.exists():
        return bundled
    return Path(__file__).resolve().parent.parent / "templates"


_jinja_env = Environment(
    loader=FileSystemLoader(str(_template_dir())),
    autoescape=select_autoescape(["html", "xml"]),
)


def _data_uri_from_b64(b64_string: str, default_mime: str = "image/png") -> str:
    # Clean data URI scheme if already present
    if b64_string.startswith("data:"):
        header_end = b64_string.find(",")
        if header_end != -1:
            b64_string = b64_string[header_end + 1 :]

    mime = default_mime
    try:
        sample = b64_string[:24]
        # Pad sample to a multiple of 4 to ensure valid base64 decoding format
        sample += "=" * ((4 - len(sample) % 4) % 4)
        header = base64.b64decode(sample, validate=True)
        if header.startswith(b"\xff\xd8"):
            mime = "image/jpeg"
        elif header.startswith(b"\x89PNG"):
            mime = "image/png"
        elif header.startswith(b"RIFF") and header[8:12] == b"WEBP":
            mime = "image/webp"
    except Exception:
        pass
    return f"data:{mime};base64,{b64_string}"


def _data_uri_from_bytes(content: bytes, default_mime: str = "image/png") -> str:
    b64 = base64.b64encode(content).decode("ascii")
    mime = default_mime
    if content.startswith(b"\xff\xd8"):
        mime = "image/jpeg"
    elif content.startswith(b"\x89PNG"):
        mime = "image/png"
    elif content.startswith(b"RIFF") and content[8:12] == b"WEBP":
        mime = "image/webp"
    return f"data:{mime};base64,{b64}"


def _serialize_image(ref: PanelImageRef) -> dict[str, Any]:
    return {
        "filename": ref.filename,
        "caption": ref.caption,
        "position": ref.position,
    }


def _serialize_panel(panel: Panel) -> dict[str, Any]:
    return {
        "cuadrante": panel.cuadrante,
        "fecha_corte": panel.fecha_corte,
        "motivo": panel.motivo,
        "imagenes": [_serialize_image(img) for img in panel.imagenes],
        "source_row_index": panel.source_row_index,
    }


def _prepare_logos(logos: dict[str, str | None]) -> tuple[str | None, str | None, str | None]:
    """Devuelve (logo_left, logo_right, logo_center) como data URIs."""
    left_raw = logos.get("left")
    right_raw = logos.get("right")
    left = _data_uri_from_b64(left_raw) if left_raw else None
    right = _data_uri_from_b64(right_raw) if right_raw else None
    center = None
    if left and not right:
        center = left
    elif right and not left:
        center = right
    elif left and right:
        # Cuando ambos logos existen, mostramos solo el izquierdo centrado
        # para coincidir con el documento Word de referencia.
        center = left
    return left, right, center


def render_pdf(
    panels: tuple[Panel, ...],
    logos: dict[str, str | None],
    images: dict[str, str],
    export_mode: ExportMode,
    image_paths: dict[str, str] | None = None,
) -> tuple[bytes, str]:
    """Renderiza un PDF consolidado con una p\u00e1gina por Panel."""
    if not panels:
        msg = "No hay paneles para exportar"
        raise RenderingError(msg)

    try:
        template = _jinja_env.get_template("panel-aviso-corte.html")
    except Exception as exc:
        logger.exception("No se pudo cargar la plantilla panel-aviso-corte.html")
        msg = f"Error al cargar plantilla: {exc}"
        raise RenderingError(msg) from exc

    logo_left, logo_right, logo_center = _prepare_logos(logos)

    image_uris: dict[str, str] = {}
    for filename, raw_path in (image_paths or {}).items():
        path = Path(raw_path)
        if path.is_file():
            image_uris[filename] = path.resolve().as_uri()
    for filename, b64 in images.items():
        image_uris.setdefault(filename, _data_uri_from_b64(b64))

    panels_data: list[dict[str, Any]] = []
    for panel in panels:
        pdict = _serialize_panel(panel)
        pdict["image_uris"] = {
            img["filename"]: image_uris.get(img["filename"])
            for img in pdict["imagenes"]
        }
        panels_data.append(pdict)

    fecha_hora = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"panel_aviso_corte_{fecha_hora}.pdf"

    context = {
        "panels": panels_data,
        "logo_left": logo_left,
        "logo_right": logo_right,
        "logo_center": logo_center,
        "fecha_hora": fecha_hora,
    }

    try:
        html_string = template.render(context)
    except Exception as exc:
        logger.exception("Error al renderizar la plantilla Jinja2")
        msg = f"Error al renderizar plantilla: {exc}"
        raise RenderingError(msg) from exc

    try:
        from weasyprint import HTML

        pdf_buffer = BytesIO()
        HTML(string=html_string, base_url=str(_template_dir())).write_pdf(pdf_buffer)
        pdf_bytes = pdf_buffer.getvalue()
    except Exception as exc:
        logger.exception("WeasyPrint fall\u00f3 al generar el PDF")
        msg = f"Error al generar PDF: {exc}"
        raise RenderingError(msg) from exc

    return pdf_bytes, filename


def render_docx(
    panels: tuple[Panel, ...],
    logos: dict[str, str | None],
    images: dict[str, str],
    export_mode: ExportMode,
    image_paths: dict[str, str] | None = None,
) -> tuple[bytes, str]:
    """Genera un documento Word (.docx) con tabla de 9 filas x 4 columnas con merges."""
    if not panels:
        msg = "No hay paneles para exportar"
        raise RenderingError(msg)

    try:
        from docx import Document
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import parse_xml
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt, RGBColor
    except ImportError as exc:
        msg = "python-docx no est\u00e1 instalado"
        raise RenderingError(msg) from exc

    doc = Document()
    doc.styles["Normal"].font.name = DOC_FONT
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT)

    def cm_to_twips(value: float) -> int:
        return round(value * 567)

    def set_row_height(row: Any, height_cm: float) -> None:
        trPr = row._tr.get_or_add_trPr()
        for node in trPr.findall(qn("w:trHeight")):
            trPr.remove(node)
        trPr.append(parse_xml(
            '<w:trHeight xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            f'w:val="{cm_to_twips(height_cm)}" w:hRule="exact"/>',
        ))

    def set_cell_width(cell: Any, width_cm: float) -> None:
        tcPr = cell._tc.get_or_add_tcPr()
        tcW = tcPr.find(qn("w:tcW"))
        if tcW is None:
            tcW = parse_xml(
                '<w:tcW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>',
            )
            tcPr.append(tcW)
        tcW.set(qn("w:w"), str(cm_to_twips(width_cm)))
        tcW.set(qn("w:type"), "dxa")

    def set_cell_margins(cell: Any, top: int = 0, left: int = CELL_MARGIN_TWIPS,
                         bottom: int = 0, right: int = CELL_MARGIN_TWIPS) -> None:
        tcPr = cell._tc.get_or_add_tcPr()
        old = tcPr.find(qn("w:tcMar"))
        if old is not None:
            tcPr.remove(old)
        tcPr.append(parse_xml(
            '<w:tcMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            f'<w:top w:w="{top}" w:type="dxa"/>'
            f'<w:start w:w="{left}" w:type="dxa"/>'
            f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
            f'<w:end w:w="{right}" w:type="dxa"/>'
            '</w:tcMar>',
        ))

    def set_vertical_align(cell: Any, value: str = "top") -> None:
        tcPr = cell._tc.get_or_add_tcPr()
        old = tcPr.find(qn("w:vAlign"))
        if old is not None:
            tcPr.remove(old)
        tcPr.append(parse_xml(
            '<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            f'w:val="{value}"/>',
        ))

    def set_no_wrap(cell: Any) -> None:
        tcPr = cell._tc.get_or_add_tcPr()
        if tcPr.find(qn("w:noWrap")) is None:
            tcPr.append(parse_xml(
                '<w:noWrap xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>',
            ))

    def format_run(run: Any, size_pt: float, *, bold: bool = False,
                   color: tuple[int, int, int] | None = None) -> None:
        run.bold = bold
        run.font.size = Pt(size_pt)
        run.font.name = DOC_FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT)
        if color:
            run.font.color.rgb = RGBColor(*color)

    def cover_image_size_cm(
        content: bytes, max_width_cm: float, max_height_cm: float,
    ) -> tuple[float, float, tuple[int, int, int, int]]:
        """Scale image to *cover* the cell and return crop offsets.

        Returns (insert_width_cm, insert_height_cm, (crop_top, crop_bottom, crop_left, crop_right))
        where crop values are in units of 1/100000 used by ``<a:srcRect>``.
        """
        from PIL import Image

        with Image.open(BytesIO(content)) as image:
            width_px, height_px = image.size

        if width_px <= 0 or height_px <= 0:
            return max_width_cm, max_height_cm, (0, 0, 0, 0)

        width_ratio = max_width_cm / width_px
        height_ratio = max_height_cm / height_px
        # Use max to *cover* (fill) the cell, opposite of *contain*
        scale = max(width_ratio, height_ratio)
        scaled_w = width_px * scale
        scaled_h = height_px * scale

        crop_left = 0
        crop_right = 0
        crop_top = 0
        crop_bottom = 0

        if scaled_w > max_width_cm + 0.01:
            excess_pct = (scaled_w - max_width_cm) / scaled_w
            half = excess_pct / 2
            crop_left = round(half * 100_000)
            crop_right = round(half * 100_000)

        if scaled_h > max_height_cm + 0.01:
            excess_pct = (scaled_h - max_height_cm) / scaled_h
            # Crop from bottom only (align top, matching preview's flex-start)
            crop_top = 0
            crop_bottom = round(excess_pct * 100_000)

        return max_width_cm, max_height_cm, (crop_top, crop_bottom, crop_left, crop_right)

    def _apply_crop(inline_shape: Any, crop: tuple[int, int, int, int]) -> None:
        """Apply Word XML cropping to an inline picture."""
        top, bottom, left, right = crop
        if top == 0 and bottom == 0 and left == 0 and right == 0:
            return
        # Navigate to the a:blipFill element in the inline shape's XML
        nsmap = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
        blip_fill = inline_shape._inline.graphic.graphicData.find(
            ".//a:blipFill", nsmap,
        )
        if blip_fill is None:
            return
        # Remove existing srcRect if any
        for old in blip_fill.findall("a:srcRect", nsmap):
            blip_fill.remove(old)
        from lxml import etree
        src_rect = etree.SubElement(
            blip_fill,
            f"{{{nsmap['a']}}}srcRect",
        )
        src_rect.set("t", str(top))
        src_rect.set("b", str(bottom))
        src_rect.set("l", str(left))
        src_rect.set("r", str(right))

    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

    image_bytes: dict[str, bytes] = {}
    for filename, b64 in images.items():
        with contextlib.suppress(Exception):
            image_bytes[filename] = base64.b64decode(b64, validate=True)

    disk_image_paths = {
        filename: Path(raw_path)
        for filename, raw_path in (image_paths or {}).items()
        if Path(raw_path).is_file()
    }

    logo_bytes: bytes | None = None
    left_raw = logos.get("left")
    right_raw = logos.get("right")
    if left_raw:
        with contextlib.suppress(Exception):
            logo_bytes = base64.b64decode(left_raw, validate=True)
    elif right_raw:
        with contextlib.suppress(Exception):
            logo_bytes = base64.b64decode(right_raw, validate=True)

    for pidx, panel in enumerate(panels):
        if pidx > 0:
            doc.add_page_break()

        table = doc.add_table(rows=9, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.allow_autofit = False

        tblPr = table._tbl.tblPr
        existing_borders = tblPr.find(qn("w:tblBorders"))
        if existing_borders is not None:
            tblPr.remove(existing_borders)
        tblPr.append(parse_xml(
            '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '</w:tblBorders>',
        ))

        tblPr.append(parse_xml(
            '<w:tblW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            f'w:w="{cm_to_twips(TABLE_WIDTH_CM)}" w:type="dxa"/>',
        ))
        grid = table._tbl.tblGrid
        for col, width_cm in zip(grid.gridCol_lst, COL_WIDTHS_CM, strict=True):
            col.set(qn("w:w"), str(cm_to_twips(width_cm)))

        widths = [Cm(width) for width in COL_WIDTHS_CM]
        for row in table.rows:
            for ci, cell in enumerate(row.cells):
                cell.width = widths[ci]
                set_cell_width(cell, COL_WIDTHS_CM[ci])
                set_cell_margins(cell)
                set_vertical_align(cell, "center")

        row_keys = (
            "title", "meta_a", "meta_b", "meta_c", "section",
            "image", "caption", "image", "caption_last",
        )
        for row, key in zip(table.rows, row_keys, strict=True):
            set_row_height(row, ROW_HEIGHTS_CM[key])

        # --- Row 0: Titulo (cols 0-2) + Logo (col 3, rowspan 4) ---
        table.cell(0, 0).merge(table.cell(0, 2))
        title_cell = table.cell(0, 0)
        title_cell.paragraphs[0].clear()
        p = title_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            "AVISO DE CORTE DEL SERVICIO DE AGUA POTABLE, "
            "POR TRABAJOS DE MEJORAMIENTO EN EL SISTEMA",
        )
        format_run(run, 12, bold=True)

        logo_cell = table.cell(0, 3).merge(table.cell(3, 3))
        logo_cell.paragraphs[0].clear()
        set_vertical_align(logo_cell, "center")
        if logo_bytes:
            p = logo_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(BytesIO(logo_bytes), width=Cm(LOGO_WIDTH_CM))

        # --- Rows 1-3: Label + Value ---
        data_items = [
            (1, "CUADRANTE AFECTADO", panel.cuadrante),
            (2, "FECHA DE CORTE", panel.fecha_corte),
            (3, "MOTIVO", panel.motivo),
        ]
        for ri, label, value in data_items:
            lbl_cell = table.cell(ri, 0)
            lbl_cell.paragraphs[0].clear()
            set_cell_margins(lbl_cell, top=40, right=0, bottom=40)
            set_no_wrap(lbl_cell)
            run = lbl_cell.paragraphs[0].add_run(label)
            format_run(run, 9, bold=True)

            # Value (merge cols 1-2)
            table.cell(ri, 1).merge(table.cell(ri, 2))
            val_cell = table.cell(ri, 1)
            val_cell.paragraphs[0].clear()
            set_cell_margins(val_cell, top=40, bottom=40)
            run = val_cell.paragraphs[0].add_run(value)
            format_run(run, 9.5, bold=True, color=TEAL_RGB)

        # --- Row 4: PANEL FOTOGRAFICO (span 4) ---
        table.cell(4, 0).merge(table.cell(4, 3))
        sec_cell = table.cell(4, 0)
        sec_cell.paragraphs[0].clear()
        p = sec_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("PANEL FOTOGRAFICO")
        format_run(run, 12, bold=True)

        # --- Rows 5-8: Imagenes y captions ---
        # (img_row, cap_row, pos_left, pos_right)
        photo_rows = [(5, 6, 1, 2), (7, 8, 3, 4)]

        for img_row, cap_row, pos_left, pos_right in photo_rows:
            # Merge: cols 0-1 y cols 2-3
            merged_cells = {
                pos_left: (
                    table.cell(img_row, 0).merge(table.cell(img_row, 1)),
                    table.cell(cap_row, 0).merge(table.cell(cap_row, 1)),
                ),
                pos_right: (
                    table.cell(img_row, 2).merge(table.cell(img_row, 3)),
                    table.cell(cap_row, 2).merge(table.cell(cap_row, 3)),
                ),
            }

            for pos, (img_cell, cap_cell) in merged_cells.items():
                img_cell.paragraphs[0].clear()
                cap_cell.paragraphs[0].clear()
                set_cell_margins(img_cell, left=0, right=0)
                set_cell_margins(cap_cell)

                img_ref = next(
                    (im for im in panel.imagenes if im.position == pos), None,
                )

                p = img_cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                image_content: bytes | None = None
                if img_ref and img_ref.filename in disk_image_paths:
                    with contextlib.suppress(Exception):
                        image_content = disk_image_paths[img_ref.filename].read_bytes()
                elif img_ref and img_ref.filename in image_bytes:
                    image_content = image_bytes[img_ref.filename]

                if image_content is not None:
                    width_cm, height_cm, crop = cover_image_size_cm(
                        image_content,
                        PHOTO_WIDTH_CM,
                        PHOTO_HEIGHT_CM,
                    )
                    run = p.add_run()
                    inline_shape = run.add_picture(
                        BytesIO(image_content),
                        width=Cm(width_cm),
                        height=Cm(height_cm),
                    )
                    _apply_crop(inline_shape, crop)
                else:
                    run = p.add_run("Sin imagen")
                    run.italic = True
                    format_run(run, 9)

                p = cap_cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                if img_ref:
                    cap_text = img_ref.caption
                else:
                    cap_text = (
                        f"IMAGEN N\u00b0{pos}: "
                        "(Indicar direcci\u00f3n seg\u00fan lista de usuarios)"
                    )
                run = p.add_run(cap_text)
                format_run(run, 12)

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    fecha_hora = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"panel_aviso_corte_{fecha_hora}.docx"
    return out.getvalue(), filename


__all__ = ["render_docx", "render_pdf"]
